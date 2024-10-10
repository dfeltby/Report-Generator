import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
from PIL import Image, ImageTk
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
import sys

# Function to recall VNA Setup File
def recall_vna_setup_file():
    """Open a dialog to select a VNA Setup file (.STA or any file)."""
    filetypes = [("VNA Setup Files", "*.STA"), ("All files", "*.*")]
    
    file_name = filedialog.askopenfilename(
        title="Select your VNA Setup File to load",
        filetypes=filetypes
    )
    
    if file_name:
        messagebox.showinfo("File Loaded", f"{file_name} loaded successfully")
    else:
        messagebox.showwarning("No File Selected", "No file was selected!")

# Function to launch the MultiStepWizard
def launch_report_wizard_wrapper():
    # Create a new window for the wizard
    wizard_window = tk.Toplevel()
    wizard_window.title("Report Generator Wizard")
    wizard_window.geometry("500x400")
    wizard = MultiStepWizard(wizard_window, {})
    wizard_window.mainloop()

# Main GUI window
def main_gui():
    root = tk.Tk()
    root.title("VNA Setup and Report Generator")
    root.configure(bg='white')

    # Load and resize the logos
    logo_path_left = r"C:/Users/davidf/OneDrive - glenairukltd.onmicrosoft.com/Documents/VNA Report Writer/glenair-logo-new.png"
    logo_path_right = r"C:/Users/davidf/OneDrive - glenairukltd.onmicrosoft.com/Documents/VNA Report Writer/Keysight P5004B.jpg"

    logo_image_left = Image.open(logo_path_left)
    logo_image_left = logo_image_left.resize((int(logo_image_left.width / 4), int(logo_image_left.height / 4)), Image.Resampling.LANCZOS)
    logo_photo_left = ImageTk.PhotoImage(logo_image_left)

    logo_image_right = Image.open(logo_path_right)
    logo_image_right = logo_image_right.resize((int(logo_image_right.width / 4), int(logo_image_right.height / 4)), Image.Resampling.LANCZOS)
    logo_photo_right = ImageTk.PhotoImage(logo_image_right)

    # Create widgets
    title_label = tk.Label(root, text="VNA Setup and Report Generator", font=("Arial", 16, "bold"), bg='blue', fg='white')
    logo_label_left = tk.Label(root, image=logo_photo_left, bg='white')
    logo_label_right = tk.Label(root, image=logo_photo_right, bg='white')

    # Place widgets in grid
    title_label.grid(row=1, column=1, padx=10, pady=10, sticky='n')
    logo_label_left.grid(row=1, column=0, padx=10, pady=10, sticky='nw')
    logo_label_right.grid(row=1, column=2, padx=(10, 20), pady=10, sticky='ne')

    # Create a frame to center the buttons
    button_frame = tk.Frame(root, bg='white')
    button_frame.grid(row=2, column=0, columnspan=3, pady=20)

    # Create buttons
    recall_vna_button = tk.Button(button_frame, text="Recall VNA Setup File", bg='white', font=("Arial", 12), padx=20, command=recall_vna_setup_file)
    generate_report_button = tk.Button(button_frame, text="Generate Report", bg='white', font=("Arial", 12), padx=20, command=launch_report_wizard_wrapper)

    # Place buttons in the frame
    recall_vna_button.pack(side='left', padx=20)
    generate_report_button.pack(side='right', padx=20)

    root.mainloop()

# MultiStepWizard class definition
class MultiStepWizard:
    def __init__(self, master, stored_values):
        self.master = master
        self.master.geometry("500x400")
        self.content_frame = tk.Frame(master)
        self.content_frame.pack(pady=20, padx=20, expand=True, fill=tk.BOTH)

        self.stored_values = stored_values
        self.selected_options = stored_values.get("selected_options", set())
        self.user_name = stored_values.get("user_name", None)
        self.job_card_number = stored_values.get("job_card", None)
        self.port_1_connector = stored_values.get("port_1_connector", None)
        self.port_2_connector = stored_values.get("port_2_connector", None)
        self.same_as_port_1 = stored_values.get("same_as_port_1", False)
        self.single_port_measurement = stored_values.get("single_port_measurement", False)
        self.tickbox_selected = stored_values.get("tickbox_selected", set())  # Tickbox state

        # Initialize calibration_data and set default values
        if "calibration_data" not in self.stored_values:
            self.stored_values["calibration_data"] = {"data1": "XNA34 Jun23-Jun25", "data2": "XRA11 Jun23-Jun25"}
        self.report_date = stored_values.get("report_date", datetime.today().strftime('%d/%m/%Y'))

        self.current_step = "step_1"
        self.job_card_entry = None

        # Updated steps order
        self.steps_order = ["step_1", "step_2", "job_card", "vna_connectors", "calibration_data", "report_date", "review"]
        self.step_index = 0

        # Button frame to keep buttons at the bottom
        self.button_frame = tk.Frame(master)
        self.button_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=10)

        self.back_button = tk.Button(self.button_frame, text="Back", command=self.back)
        self.back_button.pack(side=tk.LEFT, padx=10)

        self.next_button = tk.Button(self.button_frame, text="Next", command=self.next)
        self.next_button.pack(side=tk.RIGHT, padx=10)

        self.template_path = None  # Track the selected template path

        self.create_user_selection_step()

    def create_user_selection_step(self):
        print("[DEBUG] Creating Step 1: User Selection")
        self.clear_content_frame()
        self.current_step = "step_1"

        label = tk.Label(self.content_frame, text="Select User Template", font=("Arial", 14))
        label.pack(pady=10)

        # User options: Alexander Peet, Mark Grogan, David Feltbower
        self.user_var = tk.StringVar(value=self.stored_values.get("user_name", "Alexander Peet"))
        user_options = ["Alexander Peet", "Mark Grogan", "David Feltbower"]

        for user in user_options:
            rb = tk.Radiobutton(self.content_frame, text=user, variable=self.user_var, value=user)
            rb.pack(anchor='w')

        self.configure_buttons(back_disabled=True, next_text="Next", next_command=self.next)

    def load_user_template(self, user_name):
        template_map = {
            "Alexander Peet": "Template_AP.docx",
            "Mark Grogan": "Template_MG.docx",
            "David Feltbower": "Template_DF.docx"
        }
        template_directory = r"C:\Users\davidf\OneDrive - glenairukltd.onmicrosoft.com\Documents\VNA Report Writer"
        self.template_path = os.path.join(template_directory, template_map[user_name])
        if os.path.exists(self.template_path):
            print(f"[DEBUG] Loading template for {user_name}: {self.template_path}")
        else:
            print(f"[ERROR] Template file not found for {user_name}: {self.template_path}")

    def insert_job_card_to_template(self):
        """Insert the job card number and other selections into the Word template."""
        if self.template_path and os.path.exists(self.template_path):
            try:
                doc = Document(self.template_path)
                # Replace <Job Card p/n> in the header
                for section in doc.sections:
                    header = section.header
                    for paragraph in header.paragraphs:
                        if '<Job Card p/n>' in paragraph.text:
                            paragraph.text = paragraph.text.replace('<Job Card p/n>', self.job_card_number)
                            print(f"[DEBUG] Replaced <Job Card p/n> with {self.job_card_number}")

                # Replace placeholders in the main document body
                self.replace_placeholders_in_body(doc)

                # Save the modified document
                output_path = os.path.join(os.path.dirname(self.template_path), f"Report_{self.job_card_number}.docx")
                print(f"[DEBUG] Saving document to: {output_path}")
                doc.save(output_path)
                print(f"[DEBUG] Saved modified document as {output_path}")
            except Exception as e:
                print(f"[ERROR] Failed to modify the template or save the document: {e}")
        else:
            print(f"[ERROR] No valid template loaded or template not found at {self.template_path}")

    def replace_placeholders_in_body(self, doc):
        """Replace the placeholders in the main body of the document and set font size to 11."""
        # Ensure all values are strings or 'n/a' if None
        port_1_connector = self.stored_values.get("port_1_connector", "n/a")
        port_2_connector = self.stored_values.get("port_2_connector", "n/a")
        vna_cal = self.stored_values["calibration_data"].get("data1", "n/a")
        ecal_cal = self.stored_values["calibration_data"].get("data2", "n/a")
        report_date = self.stored_values.get("report_date", "n/a")

        # Replace placeholders in paragraphs and set font size
        for paragraph in doc.paragraphs:
            if '<Port 1>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<Port 1>', port_1_connector)
                self.set_paragraph_font_size(paragraph, 11)
                print(f"[DEBUG] Replaced <Port 1> with {port_1_connector}")
            if '<Port 2>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<Port 2>', port_2_connector)
                self.set_paragraph_font_size(paragraph, 11)
                print(f"[DEBUG] Replaced <Port 2> with {port_2_connector}")
            if '<VNA_Cal>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<VNA_Cal>', vna_cal)
                self.set_paragraph_font_size(paragraph, 11)
                print(f"[DEBUG] Replaced <VNA_Cal> with {vna_cal}")
            if '<E-Cal_Cal>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<E-Cal_Cal>', ecal_cal)
                self.set_paragraph_font_size(paragraph, 11)
                print(f"[DEBUG] Replaced <E-Cal_Cal> with {ecal_cal}")
            if '<Date>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<Date>', report_date)
                self.set_paragraph_font_size(paragraph, 11)
                print(f"[DEBUG] Replaced <Date> with {report_date}")

    def set_paragraph_font_size(self, paragraph, font_size):
        """Set the font size of all runs in the paragraph to a specified value."""
        for run in paragraph.runs:
            run.font.size = Pt(font_size)

    def create_final_step_message(self):
        """Display the final message after the wizard is complete."""
        print("[DEBUG] Displaying final message after completion")
        self.clear_content_frame()
        self.current_step = "final_message"

        label = tk.Label(self.content_frame, text="Now we are ready to start recording data from the VNA.", font=("Arial", 16, "bold"), wraplength=450)
        label.pack(pady=10)

        subtext = tk.Label(self.content_frame, text="From now on, you need to scan the GUK serial number on the Job Card and choose which derivative to save at.", font=("Arial", 12), wraplength=450)
        subtext.pack(pady=10)

        self.configure_buttons(back_disabled=True, next_text="Finish", next_command=self.master.destroy)

    def finish(self):
        """Finalize the wizard, insert data into the Word template, and save the document."""
        self.insert_job_card_to_template()  # Insert the job card number into the Word template
        
        # Show final confirmation
        messagebox.showinfo("Wizard Completed", "The report has been generated and saved successfully.")
        
        # Keep the Tkinter window alive
        self.master.mainloop()

    def create_option_selection_step(self):
        print("[DEBUG] Creating Step 2: Option Selection with Tickboxes")
        self.clear_content_frame()
        self.current_step = "step_2"

        # Main heading with text wrapping
        options_label = tk.Label(self.content_frame, text="Choose which options to place into the Report.", font=("Arial", 16, "bold"), wraplength=450)
        options_label.pack(pady=10)

        # Subtext note with wrapping
        subtext_label = tk.Label(self.content_frame, text="Note that all options are saved from the VNA, irrespective of what options are chosen here.", font=("Arial", 10), wraplength=450)
        subtext_label.pack(pady=5)

        # Options for checkboxes
        options = ["S11", "S12", "S21", "S22", "T11", "T22"]

        self.option_vars = []
        for option in options:
            var = tk.IntVar(value=1 if option in self.stored_values.get("selected_options", set()) else 0)
            cb = tk.Checkbutton(self.content_frame, text=option, variable=var,
                                command=lambda v=var, opt=option: self.toggle_option(v, opt, options))
            cb.pack(anchor='w')
            self.option_vars.append(var)

        # "Select All" checkbox
        self.select_all_var = tk.IntVar(value=1 if len(self.stored_values.get("selected_options", set())) == len(options) else 0)
        select_all_cb = tk.Checkbutton(self.content_frame, text="Select All", variable=self.select_all_var,
                                       command=lambda: self.toggle_select_all(options))
        select_all_cb.pack(anchor='w')

        self.configure_buttons(back_disabled=False, next_text="Next", next_command=self.next)

    def toggle_option(self, var, option, options):
        """Toggle individual option and manage the selected_options set."""
        if var.get() == 1:
            self.selected_options.add(option)
        else:
            self.selected_options.discard(option)
            self.select_all_var.set(0)  # Uncheck "Select All" if any option is unchecked
        self.stored_values["selected_options"] = self.selected_options

        # Check "Select All" if all options are selected
        if len(self.selected_options) == len(options):
            self.select_all_var.set(1)

    def toggle_select_all(self, options):
        """Select or Deselect all checkboxes."""
        if self.select_all_var.get() == 1:
            self.selected_options.update(options)
        else:
            self.selected_options.clear()
        self.stored_values["selected_options"] = self.selected_options

        # Update the individual checkboxes to reflect the "Select All" state
        for var in self.option_vars:
            var.set(self.select_all_var.get())

    def create_job_card_step(self):
        print("[DEBUG] Creating Step 3: Job Card Input")
        self.clear_content_frame()
        self.current_step = "job_card"

        label = tk.Label(self.content_frame, text="Enter Job Card Part Number:", font=("Arial", 14))
        label.pack(pady=10)

        self.job_card_entry = tk.Entry(self.content_frame)
        self.job_card_entry.insert(0, self.stored_values.get("job_card", ""))
        self.job_card_entry.pack(pady=5)

        self.configure_buttons(back_disabled=False, next_text="Next", next_command=self.next)

    def create_vna_connector_step(self):
        print("[DEBUG] Creating Step 4: VNA Test Port Connectors Selection")
        self.clear_content_frame()
        self.current_step = "vna_connectors"

        label = tk.Label(self.content_frame, text="Choose VNA Test Port Connectors:", font=("Arial", 14))
        label.pack(pady=10)

        # Dropdown for Port 1
        port_1_label = tk.Label(self.content_frame, text="Port 1:", font=("Arial", 12))
        port_1_label.pack(pady=5)
        self.port_1_var = tk.StringVar(value=self.stored_values.get("port_1_connector", ""))
        port_1_options = ["SMA", "N", "SMB", "SMP"]
        port_1_menu = tk.OptionMenu(self.content_frame, self.port_1_var, *port_1_options)
        port_1_menu.pack(pady=5)

        # Checkboxes for "Same as Port 1" and "Single Port Measurement"
        self.same_as_port_1_var = tk.IntVar(value=1 if self.stored_values.get("same_as_port_1", False) else 0)
        self.single_port_measurement_var = tk.IntVar(value=1 if self.stored_values.get("single_port_measurement", False) else 0)

        same_as_port_1_cb = tk.Checkbutton(self.content_frame, text="Same as Port 1", variable=self.same_as_port_1_var,
                                           command=self.toggle_same_as_port_1)
        same_as_port_1_cb.pack(pady=5)

        single_port_cb = tk.Checkbutton(self.content_frame, text="Single Port Measurement", variable=self.single_port_measurement_var,
                                        command=self.toggle_single_port_measurement)
        single_port_cb.pack(pady=5)

        # Dropdown for Port 2
        self.port_2_label = tk.Label(self.content_frame, text="Port 2:", font=("Arial", 12))
        self.port_2_label.pack(pady=5)
        self.port_2_var = tk.StringVar(value=self.stored_values.get("port_2_connector", ""))
        self.port_2_menu = tk.OptionMenu(self.content_frame, self.port_2_var, *port_1_options)
        self.port_2_menu.pack(pady=5)

        # Initially check the state of checkboxes
        self.update_port_2_visibility()

        self.configure_buttons(back_disabled=False, next_text="Next", next_command=self.next)

    def toggle_same_as_port_1(self):
        """If Same as Port 1 is selected, disable Single Port Measurement and update Port 2 behavior."""
        if self.same_as_port_1_var.get():
            self.single_port_measurement_var.set(0)  # Uncheck Single Port Measurement
            self.stored_values["same_as_port_1"] = True
            self.stored_values["single_port_measurement"] = False
        self.update_port_2_visibility()

    def toggle_single_port_measurement(self):
        """If Single Port Measurement is selected, disable Same as Port 1 and update Port 2 behavior."""
        if self.single_port_measurement_var.get():
            self.same_as_port_1_var.set(0)  # Uncheck Same as Port 1
            self.stored_values["same_as_port_1"] = False
            self.stored_values["single_port_measurement"] = True
        self.update_port_2_visibility()

    def update_port_2_visibility(self):
        """Handle visibility of Port 2 dropdown based on checkboxes."""
        if self.same_as_port_1_var.get() == 1:
            self.port_2_label.pack_forget()
            self.port_2_menu.pack_forget()
            self.port_2_var.set(self.port_1_var.get())  # Set Port 2 to same as Port 1
        elif self.single_port_measurement_var.get() == 1:
            self.port_2_label.pack_forget()
            self.port_2_menu.pack_forget()
            self.port_2_var.set("n/a")  # Set Port 2 to "n/a"
        else:
            self.port_2_label.pack(pady=5)
            self.port_2_menu.pack(pady=5)

    def create_calibration_data_step(self):
        print("[DEBUG] Creating Step 5: Select Calibration Data")
        self.clear_content_frame()
        self.current_step = "calibration_data"

        # Calibration Data 1
        label1 = tk.Label(self.content_frame, text="VNA Calibration:", font=("Arial", 14))
        label1.pack(pady=5)
        self.calibration_data_var1 = tk.StringVar(value=self.stored_values["calibration_data"].get("data1", "XNA34 Jun23-Jun25"))
        options1 = ["XNA34 Jun23-Jun25", "Add New"]
        dropdown1 = tk.OptionMenu(self.content_frame, self.calibration_data_var1, *options1, command=self.handle_add_new_vna)
        dropdown1.pack(pady=5)

        # Calibration Data 2
        label2 = tk.Label(self.content_frame, text="E-Cal Calibration:", font=("Arial", 14))
        label2.pack(pady=5)
        self.calibration_data_var2 = tk.StringVar(value=self.stored_values["calibration_data"].get("data2", "XRA11 Jun23-Jun25"))
        options2 = ["XRA11 Jun23-Jun25", "Add New"]
        dropdown2 = tk.OptionMenu(self.content_frame, self.calibration_data_var2, *options2, command=self.handle_add_new_ecal)
        dropdown2.pack(pady=5)

        self.configure_buttons(back_disabled=False, next_text="Next", next_command=self.next)

    def handle_add_new_vna(self, selection):
        if selection == "Add New":
            new_value = simpledialog.askstring("Add New VNA Calibration", "Enter new VNA Calibration value:")
            if new_value:
                self.calibration_data_var1.set(new_value)
                self.stored_values["calibration_data"]["data1"] = new_value

    def handle_add_new_ecal(self, selection):
        if selection == "Add New":
            new_value = simpledialog.askstring("Add New E-Cal Calibration", "Enter new E-Cal Calibration value:")
            if new_value:
                self.calibration_data_var2.set(new_value)
                self.stored_values["calibration_data"]["data2"] = new_value

    def create_report_date_step(self):
        print("[DEBUG] Creating Step 6: Report Date")
        self.clear_content_frame()
        self.current_step = "report_date"

        # Instruction with the date format DD/MM/YYYY
        label = tk.Label(self.content_frame, text="Enter or Confirm Report Date (format: DD/MM/YYYY):", font=("Arial", 14))
        label.pack(pady=10)

        self.date_entry = tk.Entry(self.content_frame)
        self.date_entry.insert(0, self.stored_values.get("report_date", datetime.today().strftime('%d/%m/%Y')))
        self.date_entry.pack(pady=5)

        self.configure_buttons(back_disabled=False, next_text="Next", next_command=self.next)

    def create_review_step(self):
        print("[DEBUG] Creating Review Step")
        self.clear_content_frame()
        self.current_step = "review"

        selected_options_summary = ', '.join(self.stored_values.get("selected_options", [])) if self.selected_options else 'None'
        job_card_number = self.job_card_number if self.job_card_number else 'None'
        port_1_connector = self.stored_values.get("port_1_connector", "None")
        port_2_connector = self.stored_values.get("port_2_connector", "None")
        calibration_data1 = self.stored_values["calibration_data"].get("data1", "None")
        calibration_data2 = self.stored_values["calibration_data"].get("data2", "None")
        report_date = self.stored_values.get("report_date", "None")

        # Add review page title
        title = tk.Label(self.content_frame, text="Check before submitting choices for Report", font=("Arial", 16, "bold"))
        title.pack(pady=10)

        # Display final summary
        summary = (
            f"User: {self.user_name}\n"
            f"Selected Options: {selected_options_summary}\n"
            f"Job Card: {job_card_number}\n"
            f"Port 1 Connector: {port_1_connector}\n"
            f"Port 2 Connector: {port_2_connector}\n"
            f"VNA Calibration: {calibration_data1}\n"
            f"E-Cal Calibration: {calibration_data2}\n"
            f"Report Date: {report_date}"
        )
        print(f"[DEBUG] Final Summary: \n{summary}")

        label = tk.Label(self.content_frame, text=summary, font=("Arial", 14))
        label.pack(pady=10)

        self.configure_buttons(back_disabled=False, next_text="Finish", next_command=self.finish)

    def next(self):
        # Capture data before moving forward
        if self.current_step == "step_1":
            self.user_name = self.user_var.get()
            self.stored_values["user_name"] = self.user_name
            self.load_user_template(self.user_name)
        elif self.current_step == "job_card":
            self.capture_job_card_number()
        elif self.current_step == "vna_connectors":
            # Ensure Port 1 and Port 2 values are captured correctly
            self.port_1_connector = self.port_1_var.get()
            self.port_2_connector = self.port_2_var.get()
            self.stored_values["port_1_connector"] = self.port_1_connector
            self.stored_values["port_2_connector"] = self.port_2_connector
            self.stored_values["same_as_port_1"] = self.same_as_port_1_var.get()
            self.stored_values["single_port_measurement"] = self.single_port_measurement_var.get()
        elif self.current_step == "calibration_data":
            self.stored_values["calibration_data"]["data1"] = self.calibration_data_var1.get()
            self.stored_values["calibration_data"]["data2"] = self.calibration_data_var2.get()
        elif self.current_step == "report_date":
            self.stored_values["report_date"] = self.date_entry.get()

        # Move to the next step
        self.step_index = (self.step_index + 1) % len(self.steps_order)
        self.navigate_to_step(self.steps_order[self.step_index])

    def back(self):
        # Capture data before moving back
        if self.current_step == "job_card":
            self.capture_job_card_number()
        elif self.current_step == "vna_connectors":
            self.stored_values["port_1_connector"] = self.port_1_var.get()
            self.stored_values["port_2_connector"] = self.port_2_var.get()
            self.stored_values["same_as_port_1"] = self.same_as_port_1_var.get()
            self.stored_values["single_port_measurement"] = self.single_port_measurement_var.get()
        elif self.current_step == "calibration_data":
            self.stored_values["calibration_data"]["data1"] = self.calibration_data_var1.get()
            self.stored_values["calibration_data"]["data2"] = self.calibration_data_var2.get()
        elif self.current_step == "report_date":
            self.stored_values["report_date"] = self.date_entry.get()

        # Move to the previous step
        self.step_index = (self.step_index - 1) % len(self.steps_order)
        self.navigate_to_step(self.steps_order[self.step_index])

    def navigate_to_step(self, step):
        if step == "step_1":
            self.create_user_selection_step()
        elif step == "step_2":
            self.create_option_selection_step()
        elif step == "job_card":
            self.create_job_card_step()
        elif step == "vna_connectors":
            self.create_vna_connector_step()
        elif step == "calibration_data":
            self.create_calibration_data_step()
        elif step == "report_date":
            self.create_report_date_step()
        elif step == "review":
            self.create_review_step()

    def finish(self):
        """Finalize the wizard, insert data into the Word template, and save the document."""
        self.insert_job_card_to_template()  # Insert the job card number into the Word template
        
        # Show final confirmation
        messagebox.showinfo("Wizard Completed", "The report has been generated and saved successfully.")
        
        # Keep the Tkinter window alive
        self.master.mainloop()

    def clear_content_frame(self):
        print(f"[DEBUG] Clearing content frame. Current step: {self.current_step}")
        for widget in self.content_frame.winfo_children():
            print(f"[DEBUG] Destroying widget: {widget}")
            widget.destroy()

    def configure_buttons(self, back_disabled, next_text, next_command):
        """Helper method to configure back and next buttons."""
        self.back_button.config(state=tk.DISABLED if back_disabled else tk.NORMAL)
        self.next_button.config(state=tk.NORMAL, text=next_text, command=next_command)

    def capture_job_card_number(self):
        """Capture job card number before leaving job card step."""
        if self.job_card_entry is not None and self.job_card_entry.winfo_exists():
            self.job_card_number = self.job_card_entry.get()
            self.stored_values["job_card"] = self.job_card_number
            print(f"[DEBUG] Captured job card number: {self.job_card_number}")

# Call the main GUI function
if __name__ == "__main__":
    main_gui()